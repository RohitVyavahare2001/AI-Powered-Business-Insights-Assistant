from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any
import io

class ReportGenerator:
    def __init__(self):
        """Initialize the report generator with default styling options."""
        self.default_font = "Calibri"
        self.default_font_size = Pt(11)
        self.heading_font_size = Pt(14)
    
    def create_report(self, analysis_type: str, query: str, insights: Dict[str, Any]) -> bytes:
        """
        Create a formatted report document from the insights.
        
        Args:
            analysis_type (str): Type of analysis performed
            query (str): Original business query
            insights (Dict[str, Any]): Generated insights and analysis
            
        Returns:
            bytes: The generated report as bytes for download
        """
        doc = Document()
        
        # Set document style
        self._set_document_style(doc)
        
        # Add title
        title = doc.add_heading(f"Business Insights Report: {analysis_type}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add query section
        doc.add_heading("Business Query", 1)
        doc.add_paragraph(query)
        
        # Add executive summary
        if insights.get("summary"):
            doc.add_heading("Executive Summary", 1)
            doc.add_paragraph(insights["summary"])
        
        # Add detailed analysis
        if insights.get("detailed_analysis"):
            doc.add_heading("Detailed Analysis", 1)
            doc.add_paragraph(insights["detailed_analysis"])
        
        # Add key insights and recommendations
        if insights.get("key_insights"):
            doc.add_heading("Key Insights and Recommendations", 1)
            doc.add_paragraph(insights["key_insights"])
        
        # Add action items
        if insights.get("action_items"):
            doc.add_heading("Action Items", 1)
            doc.add_paragraph(insights["action_items"])
        
        # Add risks and mitigation strategies
        if insights.get("risks"):
            doc.add_heading("Risks and Mitigation Strategies", 1)
            doc.add_paragraph(insights["risks"])
        
        # Save document to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        return doc_bytes.getvalue()
    
    def _set_document_style(self, doc: Document) -> None:
        """
        Set the default document style.
        
        Args:
            doc (Document): The document to style
        """
        # Set default font for the document
        style = doc.styles['Normal']
        style.font.name = self.default_font
        style.font.size = self.default_font_size
        
        # Style headings
        for i in range(1, 4):
            heading_style = doc.styles[f'Heading {i}']
            heading_style.font.name = self.default_font
            heading_style.font.size = self.heading_font_size
            heading_style.font.bold = True 